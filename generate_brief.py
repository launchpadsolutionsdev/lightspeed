"""
Generates Lightspeed Technical Brief (Word .docx) for senior developer / M&A review.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
PINK       = RGBColor(0xE8, 0x40, 0x7A)   # #E8407A – Lightspeed pink
CORAL      = RGBColor(0xF0, 0x68, 0x50)   # #F06850 – coral mid
GOLD       = RGBColor(0xF5, 0xA6, 0x23)   # #F5A623 – warm gold
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
MID_GREY   = RGBColor(0x55, 0x55, 0x66)   # body grey
LIGHT_BG   = RGBColor(0xFA, 0xF8, 0xF8)   # near-white tint
RULE_PINK  = RGBColor(0xE8, 0x40, 0x7A)   # horizontal rule


def set_font(run, name="Calibri", size=10, bold=False, italic=False, colour=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour


def add_paragraph(doc, text="", style="Normal", space_before=0, space_after=4,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    if text:
        run = p.add_run(text)
    return p


def add_horizontal_rule(doc, colour=RULE_PINK, thickness=8):
    """Insert a coloured top-border paragraph as a visual rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(thickness))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), str(colour))
    pBdr.append(top)
    pPr.append(pBdr)


def shade_cell(cell, hex_colour="FAF8F8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    val.get("sz",  "4"))
            el.set(qn("w:space"), val.get("space", "0"))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBdr.append(el)
    tcPr.append(tcBdr)


def add_gradient_bar(doc):
    """Simulate the pink→coral→gold gradient as a 3-column borderless table."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(2.5), Inches(2.5), Inches(2.5)]
    colours = ["E8407A", "F06850", "F5A623"]
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = widths[i]
        shade_cell(cell, colours[i])
        cell.height = Emu(90720)   # ~0.1 inch
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        # Remove borders
        set_cell_border(cell,
            top={"val":"nil"}, bottom={"val":"nil"},
            left={"val":"nil"}, right={"val":"nil"})
    return tbl


# ── Tool data ────────────────────────────────────────────────────────────────
TOOLS = [
    {
        "number": "01",
        "name":   "Response Assistant",
        "desc":   (
            "AI-powered customer support composer for lottery inquiries. "
            "Staff paste a customer question; the tool returns a complete, "
            "brand-consistent response respecting tone, language, format "
            "(email or Facebook), and compliance rules."
        ),
        "how":    (
            "Stack: Node.js/Express · PostgreSQL · Anthropic Claude Sonnet 4.6 · SSE streaming. "
            "Pattern: multi-stage RAG pipeline. "
            "Flow: PostgreSQL full-text search returns ≤30 KB candidates → "
            "Claude Haiku 4.5 filters to 8 most relevant → rated-example retrieval "
            "(30+ positive / 15 negative from response history) → Haiku topic-filters "
            "to 8+5 → corrections pipeline (all negative-rated with staff feedback → "
            "Haiku → ≤5 highest-priority overrides) → server-side system prompt assembly "
            "→ Sonnet 4.6 with Anthropic Prompt Caching (ephemeral) → SSE stream to client."
        ),
        "features": [
            "Three-tier RAG: KB entries, few-shot rated examples, staff correction overrides",
            "Anthropic Prompt Caching — reduces latency and token cost on repeated org context",
            "Server-side prompt assembly — prompt never exposed in browser DevTools",
            "Prompt-injection sanitisation with XML-wrapped user content",
            "Real-time streaming with KB citation attribution",
            "Thumbs-up/down rating → auto-creates correction KB entries for future calls",
            "Per-format token budgets (200 for Facebook, 1024 for email)",
        ],
        "ip": (
            "The tiered corrections pipeline — treating negatively-rated responses with staff "
            "feedback as highest-priority in-context overrides — is a proprietary refinement "
            "loop that continuously improves output quality without fine-tuning or retraining."
        ),
    },
    {
        "number": "02",
        "name":   "Draft Assistant",
        "desc":   (
            "Content generation tool for marketing emails, social posts, donation appeals, "
            "and event announcements. Produces brand-consistent copy informed by the "
            "organisation's internal knowledge base and optional Shopify product catalogue."
        ),
        "how":    (
            "Stack: Node.js/Express · PostgreSQL · Anthropic Claude Sonnet 4.6 · Shopify API (optional). "
            "Pattern: prompt augmentation with internal KB and e-commerce context. "
            "Flow: brand voice fetched from org profile → all internal KB entries retrieved "
            "(kb_type='internal') → Haiku relevance-filters to 8 entries → optional Shopify "
            "product catalogue (≤15 products with name/price/description) injected → "
            "non-streaming Claude Sonnet call returns complete draft."
        ),
        "features": [
            "Internal-only KB separation (kb_type='internal') keeps operational data out of support responses",
            "Shopify product context injection for e-commerce marketing copy",
            "Configurable tone (warm/professional/casual) and length (short/medium/long)",
            "Draft types: email, social post, announcement, or freeform",
        ],
        "ip": (
            "KB-type partitioning (support vs. internal) lets the same knowledge store power "
            "two distinct tools with zero cross-contamination — a deliberate architectural "
            "separation with IP value in the multi-tool SaaS model."
        ),
    },
    {
        "number": "03",
        "name":   "Ask Lightspeed",
        "desc":   (
            "Full general-purpose AI chat assistant embedded in the platform — the organisation's "
            "own Claude. Supports multi-turn conversation, file analysis, model selection, and "
            "a Teach Mode that captures organisational knowledge directly into the shared KB."
        ),
        "how":    (
            "Stack: Node.js/Express · PostgreSQL (JSONB message store) · Anthropic Claude Sonnet/Opus 4.6 "
            "· Haiku 4.5 (titles, summarisation) · SSE streaming. "
            "Pattern: stateful multi-turn conversation with dynamic context compression. "
            "Flow: full conversation history maintained in JSONB → when thread grows long, "
            "Haiku summarises all but the last 6 messages → draw schedule + rated examples "
            "injected into system prompt → user message (text or vision blocks) appended → "
            "Sonnet or Opus streams response → Teach Mode detection triggers optional KB "
            "write → turn saved to response_history for cross-tool few-shot use."
        ),
        "features": [
            "Model switchable per-session: Sonnet 4.6 (default) or Opus 4.6",
            "Automatic conversation summarisation keeps token count within model limits",
            "File uploads: images (vision), PDFs, and text with inline rendering",
            "Teach Mode: NLP pattern detection auto-offers KB save; no admin UI required",
            "Team view: colleagues' conversations visible within org scope",
            "Conversation search, archiving, and shareable prompt library",
        ],
        "ip": (
            "Teach Mode — detecting natural-language knowledge declarations and routing them "
            "into the shared KB — creates a passive, conversation-driven knowledge accumulation "
            "loop that improves all tools over time without explicit data-entry workflows."
        ),
    },
    {
        "number": "04",
        "name":   "Insights Engine",
        "desc":   (
            "Data analysis tool that accepts uploaded spreadsheets (CSV/Excel/JSON) or pulls "
            "live Shopify analytics, then returns structured business insights, trend analysis, "
            "and actionable recommendations."
        ),
        "how":    (
            "Stack: Node.js/Express · PostgreSQL · Anthropic Claude Sonnet 4.6 · Shopify API (optional). "
            "Pattern: report-type-driven prompt templates with full data injection. "
            "Flow: report type selected (customer purchases / sellers / payment tickets / Shopify / custom) "
            "→ corresponding prompt template chosen server-side → organisation brand voice optionally prepended "
            "→ full data JSON-stringified into the user message → single non-streaming Sonnet call returns "
            "structured analysis. Shopify path fetches live store data directly without requiring file upload."
        ),
        "features": [
            "Five report-type prompt templates targeting nonprofit lottery use cases",
            "Direct Shopify analytics pull (no file export required)",
            "Structured output: key metrics, trend identification, ranked recommendations",
            "Brand-voice-aware narrative framing",
        ],
        "ip": (
            "Template-switched prompt architecture maps cleanly to the domain model "
            "(nonprofit lottery operations) — easily extensible to new report types "
            "without infrastructure changes."
        ),
    },
    {
        "number": "05",
        "name":   "List Normalizer",
        "desc":   (
            "Spreadsheet cleaning and transformation tool. Users upload or paste data and "
            "describe their transformation in plain English; the tool cleans, deduplicates, "
            "and optionally generates executable JavaScript to run the transformation client-side."
        ),
        "how":    (
            "Stack: Node.js/Express · Anthropic Claude Sonnet 4.6 · browser JS runtime. "
            "Pattern: code generation with client-side execution sandbox. "
            "Flow: data + natural-language instructions sent to Sonnet → in Transform Mode, "
            "model returns only a raw JS function body (no markdown, no imports, no async) → "
            "frontend wraps it in new Function() and executes against every row in-browser → "
            "deduplication applied → CSV exported → usage logged separately (no tokens charged "
            "for client-side execution). JSON Mode returns transformed array directly. "
            "No data leaves the browser after the initial Claude call."
        ),
        "features": [
            "Transform Mode: AI-generated JS function runs client-side — no second server round-trip",
            "JSON Mode: direct structured output for programmatic downstream use",
            "Deduplication handled separately from transformation logic",
            "Usage split: tokens charged once for generation; client-side execution billed at zero",
            "Strict output contract enforced via system prompt (function body only, plain JS)",
        ],
        "ip": (
            "Offloading row-by-row execution to the browser after a single AI function-generation "
            "call is a cost-efficient architecture that scales to large datasets without "
            "server-side compute or per-row API calls."
        ),
    },
    {
        "number": "06",
        "name":   "Rules of Play Generator",
        "desc":   (
            "Legal document generator that produces complete, submission-ready Rules of Play "
            "for charitable lotteries (50/50, Catch The Ace, Prize Raffle, House Lottery) "
            "compliant with Canadian provincial gaming regulations."
        ),
        "how":    (
            "Stack: Node.js/Express · PostgreSQL · Anthropic Claude Sonnet 4.6 · Mammoth.js (DOCX parsing). "
            "Pattern: layered static-context injection with optional reference document RAG. "
            "Flow: user selects raffle type → raffle-type template (TEMPLATE_5050 etc.) chosen from "
            "ropTemplates.js → jurisdiction row fetched from DB (regulatory body, minimum age, "
            "geographic restrictions, unclaimed prize rules) → form data (org name, licence number, "
            "draw schedule, pricing tiers, prize details) merged → optional uploaded reference "
            "document text extracted via Mammoth and appended → buildSystemPrompt() assembles "
            "all four layers → single Sonnet call returns complete legal document → saved to DB "
            "→ exported as .doc (HTML-wrapped, Content-Type: application/msword)."
        ),
        "features": [
            "Four raffle-type document templates covering all common Canadian charitable lottery formats",
            "Jurisdictions table stores regulatory data independently — adding a province requires no code change",
            "Reference document upload provides style/precedent guidance without hard-coding examples",
            "Draft persistence: create, edit, regenerate, delete with full audit trail",
            "One-click .doc export with auto-detected heading and list formatting",
        ],
        "ip": (
            "The jurisdictions table as a first-class data entity decouples regulatory knowledge "
            "from application code — a defensible data asset that grows with each new "
            "province added, creating a regulatory compliance knowledge moat."
        ),
    },
]


# ── Document assembly ─────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page margins — narrow to maximise content on one page
    for section in doc.sections:
        section.top_margin    = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # ── Gradient accent bar ──────────────────────────────────────────────────
    add_gradient_bar(doc)

    # ── Confidentiality banner ───────────────────────────────────────────────
    banner = doc.add_paragraph()
    banner.paragraph_format.space_before = Pt(5)
    banner.paragraph_format.space_after  = Pt(2)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = banner.add_run("CONFIDENTIAL  —  UNDER NDA  —  FOR AUTHORISED RECIPIENTS ONLY")
    set_font(br, size=7.5, bold=True, colour=RGBColor(0xC0, 0x20, 0x50))

    # ── Title block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(3)
    title_p.paragraph_format.space_after  = Pt(1)
    tr = title_p.add_run("Lightspeed AI Platform")
    set_font(tr, size=22, bold=True, colour=PINK)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(2)
    sr = sub_p.add_run("Technical Brief  ·  Senior Developer Review  ·  March 2026")
    set_font(sr, size=9.5, italic=False, colour=MID_GREY)

    # Platform overview line
    ov_p = doc.add_paragraph()
    ov_p.paragraph_format.space_before = Pt(0)
    ov_p.paragraph_format.space_after  = Pt(5)
    ovr = ov_p.add_run(
        "Full-stack AI productivity suite for Canadian charitable lottery nonprofits. "
        "Six tools on a shared Node.js/Express + PostgreSQL backend, all powered by Anthropic Claude "
        "(Sonnet 4.6 primary · Haiku 4.5 for filtering/summarisation · Opus 4.6 optional). "
        "Single-page vanilla JS frontend. SSE streaming throughout. "
        "JWT auth · Stripe billing · Shopify integration."
    )
    set_font(ovr, size=8.5, colour=DARK_TEXT)

    add_horizontal_rule(doc)

    # ── Tool sections ────────────────────────────────────────────────────────
    for i, tool in enumerate(TOOLS):
        # Tool header row (number + name)
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(5 if i > 0 else 3)
        hdr.paragraph_format.space_after  = Pt(1)

        num_run = hdr.add_run(f"  {tool['number']}  ")
        set_font(num_run, size=9, bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF))
        # Background colour hack via character shading is not natively supported;
        # instead we render the number inline with a gradient colour.
        num_run.font.color.rgb = CORAL

        name_run = hdr.add_run(tool["name"].upper())
        set_font(name_run, size=10.5, bold=True, colour=DARK_TEXT)

        # Description
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_before = Pt(1)
        desc_p.paragraph_format.space_after  = Pt(2)
        desc_r = desc_p.add_run(tool["desc"])
        set_font(desc_r, size=8.5, colour=DARK_TEXT)

        # Two-column table: How It Works | Features + IP
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        left_cell  = tbl.rows[0].cells[0]
        right_cell = tbl.rows[0].cells[1]

        # Column widths
        left_cell._tc.get_or_add_tcPr().append(
            OxmlElement("w:tcW")
        )
        # Set widths via table layout
        tbl.columns[0].width = Inches(3.55)
        tbl.columns[1].width = Inches(3.55)

        shade_cell(left_cell,  "FAF4F7")
        shade_cell(right_cell, "FFF8F0")

        for cell in (left_cell, right_cell):
            set_cell_border(cell,
                top={"val":"single","sz":"4","color":"E8407A"},
                bottom={"val":"nil"}, left={"val":"nil"}, right={"val":"nil"})

        # Left — How It Works
        lp = left_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(2)
        lh = lp.add_run("HOW IT WORKS")
        set_font(lh, size=7, bold=True, colour=PINK)

        lp2 = left_cell.add_paragraph()
        lp2.paragraph_format.space_before = Pt(1)
        lp2.paragraph_format.space_after  = Pt(4)
        lp2.paragraph_format.left_indent  = Inches(0.05)
        lr2 = lp2.add_run(tool["how"])
        set_font(lr2, size=7.8, colour=MID_GREY)

        # Right — Key Features
        rp = right_cell.paragraphs[0]
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after  = Pt(1)
        rh = rp.add_run("KEY FEATURES")
        set_font(rh, size=7, bold=True, colour=CORAL)

        for feat in tool["features"]:
            fp = right_cell.add_paragraph()
            fp.paragraph_format.space_before = Pt(0)
            fp.paragraph_format.space_after  = Pt(1)
            fp.paragraph_format.left_indent  = Inches(0.1)
            fr = fp.add_run(f"→  {feat}")
            set_font(fr, size=7.5, colour=DARK_TEXT)

        # IP callout
        ip_p = right_cell.add_paragraph()
        ip_p.paragraph_format.space_before = Pt(3)
        ip_p.paragraph_format.space_after  = Pt(4)
        ip_p.paragraph_format.left_indent  = Inches(0.05)
        ip_label = ip_p.add_run("IP NOTE  ")
        set_font(ip_label, size=7, bold=True, colour=GOLD)
        ip_text = ip_p.add_run(tool["ip"])
        set_font(ip_text, size=7.5, italic=True, colour=MID_GREY)

        # Divider between tools (except after last)
        if i < len(TOOLS) - 1:
            add_horizontal_rule(doc, colour=RGBColor(0xE8, 0xD0, 0xD8), thickness=4)

    # ── Footer ───────────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(3)
    foot.paragraph_format.space_after  = Pt(0)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "Lightspeed  ·  Launchpad Solutions  ·  Confidential & Proprietary  ·  "
        "Distribution restricted to authorised parties under executed NDA"
    )
    set_font(fr, size=7, colour=MID_GREY)

    out = "/home/user/lightspeed/Lightspeed_Technical_Brief.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_doc()
